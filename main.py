from bs4 import BeautifulSoup
from docx import Document
class BranchStruct:
    def __init__(self, branch_node, leaf_node_list):
        self.branch_node = branch_node#主节点下面的分支节点，描述更改的项目
        self.leaf_node_list = leaf_node_list#分支节点下的叶子节点列表，描述所有的更改

class NodeStruct:
    def __init__(self,main_node, branch):
        self.main_node = main_node#主节点
        self.branch = branch#分支列表，列表中每个元素都是一个Branch

soup = BeautifulSoup(open('W2150(H-I).html', encoding='utf-8'), features="html.parser")
D1 = soup.body.contents[1].string.replace("\n","").replace(" ","")
D2 = soup.body.contents[2].string.replace("\n","").replace(" ","")
D2 = soup.body.contents[2].string.replace("\n","").replace(" ","")
body_table_node = soup.body.contents[4]
body_table_node = body_table_node.contents[3:]#前三个元素都是无效元素，体现的是D1和D2的列表信息
node_list = []

def find_node_index(searchlist ,level):
    result = []
    step = 2
    for i in range(0, len(searchlist), step):
        if searchlist[i].td.div["style"] == "margin-left: " + str(level) + "em;":
            result.append(i)
    return result

mainnode_index = find_node_index(body_table_node, level=2)#返回2级主节点的索引
for i in range(0, len(mainnode_index)-1, 1):
    main_node = body_table_node[mainnode_index[i]]#主节点
    branches = []#分支列表，每个元素都是一个Branch
    branch_node = []#分支节点
    leaf_node_list = []#叶子节点列表
    if mainnode_index[i]+2 == mainnode_index[i+1]:#说明这个主节点下2图纸没有差异，中间只是隔了一个换行符
        branches.append(None)
    else:
        branchnode_list = body_table_node[mainnode_index[i]+2: mainnode_index[i + 1]]
        branchnode_index = find_node_index(branchnode_list, level=3)
        if len(branchnode_index) == 1:#只有一个树枝
            branch_node = branchnode_list[branchnode_index[0]]
            leafnode_list = branchnode_list[branchnode_index[0] + 2: ]
            leafnode_index = find_node_index(leafnode_list, level=4)
            for k in range(0, len(leafnode_index), 1):
                leaf_node_list.append(leafnode_list[leafnode_index[k]])
            branches.append(BranchStruct(branch_node, leaf_node_list))#生成树枝
        else:#有多根树枝
            for j in range(0, len(branchnode_index) - 1, 1):
                branch_node = branchnode_list[branchnode_index[j]]
                leafnode_list = branchnode_list[branchnode_index[j] + 2: branchnode_index[j + 1]]
                leafnode_index = find_node_index(leafnode_list, level=4)
                for k in range(0, len(leafnode_index), 1):
                    leaf_node_list.append(leafnode_list[leafnode_index[k]])
                branches.append(BranchStruct(branch_node, leafnode_list))

    node = NodeStruct(main_node, branches)
    node_list.append(node)



# for i in range(3,len(body_table_node.contents),2):
#     if body_table_node.contents[i].td.div["style"] == "margin-left: 2em;":#找到主节点
#         main_node = body_table_node.contents[i]
#         item_node = []
#         leaf_node = []
#         for j in range(i+2, len(body_table_node.contents),2):
#             if body_table_node.contents[j].td.div["style"] != "margin-left: 2em;":
#
#                 if body_table_node.contents[j].td.div["style"] == "margin-left: 3em;":
#                     item_node.append(body_table_node.contents[j])
#                     for k in range(j+2,len(body_table_node.contents),2):
#                         if body_table_node.contents[j].td.div["style"] == "margin-left: 4em;":
#                             leaf_node = body_table_node.contents[k]
#                         else:
#                             break
#
#         node = NodeStruct(main_node, item_node,leaf_node)
#         node_list.append(node)

                # node_struct[node_index].item_node.append(None)#2个主节点之间是空的，表示当前主节点无差异
                # node_struct[node_index].leaf_node.append(None)



document = Document()
document.add_paragraph(D1)
document.add_paragraph(D2)
document.save("test.docx")
